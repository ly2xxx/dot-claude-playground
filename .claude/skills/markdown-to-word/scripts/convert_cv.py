import re
import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0056B3')
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    f = OxmlElement('w:rFonts')
    f.set(qn('w:ascii'), 'Calibri')
    f.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(f)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '19')  # ~9.5 pt
    rPr.append(sz)

    new_run.append(rPr)
    new_run_text = OxmlElement('w:t')
    new_run_text.text = text
    new_run.append(new_run_text)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_paragraph_bottom_border(paragraph, color_hex="1A365D", size="8"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)


def add_formatted_text(paragraph, text, font_size=9.5, font_name="Calibri", color=None,
                       default_bold=False, default_italic=False):
    """Parse inline **bold** and [text](url) markdown, append styled runs to paragraph."""
    pattern = re.compile(r'(\*\*.*?\*\*|\[.*?\]\(.*?\))')
    parts = pattern.split(text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.italic = default_italic
            run.font.name = font_name
            run.font.size = Pt(font_size)
            if color:
                run.font.color.rgb = color
        elif part.startswith('[') and '](' in part and part.endswith(')'):
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if link_match:
                add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        else:
            run = paragraph.add_run(part)
            run.bold = default_bold
            run.italic = default_italic
            run.font.name = font_name
            run.font.size = Pt(font_size)
            if color:
                run.font.color.rgb = color


def _is_contact_line(line):
    """A header line carrying an email or a markdown link is the contact line,
    not the tagline — even when the whole line is wrapped in **bold**."""
    return '@' in line or '](' in line or 'http' in line


def apply_paragraph_format(p, space_before=0, space_after=3, line_spacing=1.15):
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing


def add_skill_backtick_bullet(doc, label, content):
    """Render a `Label`  content (or **Label:** content) skill line as a bullet
    with a bolded label prefix."""
    p = doc.add_paragraph(style='List Bullet')
    apply_paragraph_format(p, space_before=0, space_after=2, line_spacing=1.15)

    run = p.add_run(f"{label}:  ")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(51, 51, 51)

    add_formatted_text(p, content, font_size=9.5, color=RGBColor(51, 51, 51))


def add_italic_note(doc, text, space_before=2, space_after=4):
    """Muted italic lead-in line (section intros, degree/timeframe lines)."""
    p = doc.add_paragraph()
    apply_paragraph_format(p, space_before=space_before, space_after=space_after,
                           line_spacing=1.15)
    add_formatted_text(p, text, font_size=9.5, color=RGBColor(85, 85, 85),
                       default_italic=True)
    return p


def add_body_bullet(doc, text, space_after=2.5):
    p = doc.add_paragraph(style='List Bullet')
    apply_paragraph_format(p, space_before=0, space_after=space_after, line_spacing=1.15)
    add_formatted_text(p, text, font_size=9.5, color=RGBColor(51, 51, 51))
    return p


def add_subheader(doc, text, size=10, color=RGBColor(17, 17, 17),
                  space_before=6, space_after=2):
    p = doc.add_paragraph()
    apply_paragraph_format(p, space_before=space_before, space_after=space_after,
                           line_spacing=1.0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def main():
    if len(sys.argv) > 1:
        md_path = os.path.abspath(sys.argv[1])
    else:
        print("Usage: python convert_cv.py <path_to_markdown_file>")
        return

    if not os.path.exists(md_path):
        print(f"Error: File '{md_path}' not found.")
        return

    docx_path = os.path.splitext(md_path)[0] + ".docx"

    print(f"Reading: {md_path}")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Page margins: Top/Bottom 0.6in, Left/Right 0.7in (matches HarveyNash)
    for section in doc.sections:
        section.top_margin = Inches(0.60)
        section.bottom_margin = Inches(0.60)
        section.left_margin = Inches(0.70)
        section.right_margin = Inches(0.70)

    current_section = "HEADER"
    # Backtick skill line pattern: `Label`  content
    backtick_re = re.compile(r'^`(\w[\w\s]*)`\s+(.+)$')
    # Bold-label skill line pattern: **Label:** content
    boldlabel_re = re.compile(r'^\*\*([^*]+?):\*\*\s+(.+)$')
    # Raw HTML block lines (e.g. <div align="center" markdown="1"> ... </div>).
    # These are GitHub-rendering hints only; Word layout is driven by the styling
    # rules below, so they must be skipped rather than printed literally.
    html_line_re = re.compile(r'^</?(div|p|span|br|center|img|a)\b[^>]*>$', re.I)

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if html_line_re.match(line):
            i += 1
            continue

        # H2 → section header with bottom border
        if line.startswith('## '):
            sec_title = line[3:].strip()
            current_section = sec_title.upper()

            p = doc.add_paragraph()
            apply_paragraph_format(p, space_before=11, space_after=4, line_spacing=1.0)
            p.paragraph_format.keep_with_next = True

            run = p.add_run(sec_title.upper())
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(26, 54, 93)   # #1A365D slate blue
            add_paragraph_bottom_border(p, color_hex="1A365D", size="8")
            i += 1
            continue

        # Skip horizontal rules
        if line == '---':
            i += 1
            continue

        # ── HEADER ──────────────────────────────────────────────────────────
        if current_section == "HEADER":
            if line.startswith('# '):
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                apply_paragraph_format(p, space_before=0, space_after=2, line_spacing=1.0)
                run = p.add_run(line[2:].strip())
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(26, 54, 93)
            elif (line.startswith('**') and line.endswith('**') and len(line) > 4
                    and not _is_contact_line(line)):
                # Tagline subtitle
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                apply_paragraph_format(p, space_before=0, space_after=4, line_spacing=1.0)
                run = p.add_run(line[2:-2].strip())
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(44, 95, 138)  # #2C5F8A accent blue
            else:
                # Contact / links line. Strip a whole-line ** wrapper first —
                # otherwise the inline parser treats the entire line as one bold
                # run and any [text](url) inside it renders as literal markdown.
                # The wrapper still carries intent, so re-apply it as default_bold.
                text = line
                emphasised = (text.startswith('**') and text.endswith('**')
                              and len(text) > 4)
                if emphasised:
                    text = text[2:-2].strip()
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                apply_paragraph_format(p, space_before=0, space_after=8, line_spacing=1.0)
                add_formatted_text(p, text, font_size=9.5, color=RGBColor(51, 51, 51),
                                   default_bold=emphasised)

        # ── PROFILE ──────────────────────────────────────────────────────────
        elif "PROFILE" in current_section:
            p = doc.add_paragraph()
            apply_paragraph_format(p, space_before=2, space_after=4, line_spacing=1.15)
            if line.startswith('**'):
                # Bold-prefixed lines (Areas of Strength, Specialised in, etc.) in accent blue
                add_formatted_text(p, line, font_size=9.5, color=RGBColor(26, 54, 93))
            else:
                add_formatted_text(p, line, font_size=9.5, color=RGBColor(51, 51, 51))

        # ── TECHNICAL SKILLS ─────────────────────────────────────────────────
        elif "SKILLS" in current_section:
            bt_match = backtick_re.match(line)
            bl_match = boldlabel_re.match(line)
            if bt_match:
                # `Core`/`Competent`/`Learning` lines → formatted bullet
                add_skill_backtick_bullet(doc, bt_match.group(1), bt_match.group(2).strip())
            elif bl_match:
                # **Proficient:** / **Familiar:** lines → same bullet treatment
                add_skill_backtick_bullet(doc, bl_match.group(1), bl_match.group(2).strip())
            elif line.startswith('**') and line.endswith('**') and len(line) > 4:
                # Skill category subheader (10pt, near-black)
                p = doc.add_paragraph()
                apply_paragraph_format(p, space_before=6, space_after=2, line_spacing=1.0)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(line[2:-2].strip())
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(17, 17, 17)
            elif line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                apply_paragraph_format(p, space_before=0, space_after=2, line_spacing=1.15)
                add_formatted_text(p, line[2:].strip(), font_size=9.5, color=RGBColor(51, 51, 51))

        # ── WORK EXPERIENCE ───────────────────────────────────────────────────
        elif "EXPERIENCE" in current_section:
            if line.startswith('### '):
                # Company/org header (10.5pt, near-black, matches HarveyNash 133350 EMU)
                p = doc.add_paragraph()
                apply_paragraph_format(p, space_before=8, space_after=2, line_spacing=1.0)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(line[4:].strip())
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(17, 17, 17)
            elif line.startswith('**') and line.endswith('**') and len(line) > 4:
                # Role / dates line (10pt, muted gray, matches HarveyNash 127000 EMU / 555555)
                p = doc.add_paragraph()
                apply_paragraph_format(p, space_before=2, space_after=2, line_spacing=1.0)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(line[2:-2].strip())
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(85, 85, 85)
            elif line.startswith('- '):
                # Bullet point
                p = doc.add_paragraph(style='List Bullet')
                apply_paragraph_format(p, space_before=0, space_after=2.5, line_spacing=1.15)
                add_formatted_text(p, line[2:].strip(), font_size=9.5, color=RGBColor(51, 51, 51))
            else:
                # One-sentence role brief (italic, muted gray)
                p = doc.add_paragraph()
                apply_paragraph_format(p, space_before=1, space_after=3, line_spacing=1.15)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(line)
                run.italic = True
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(85, 85, 85)

        # ── EDUCATION ─────────────────────────────────────────────────────────
        elif "EDUCATION" in current_section:
            if line.startswith('**') and line.endswith('**') and len(line) > 4:
                # University name subheader
                p = doc.add_paragraph()
                apply_paragraph_format(p, space_before=8, space_after=2, line_spacing=1.0)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(line[2:-2].strip())
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(17, 17, 17)
            elif line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                apply_paragraph_format(p, space_before=0, space_after=2, line_spacing=1.15)
                add_formatted_text(p, line[2:].strip(), font_size=9.5, color=RGBColor(51, 51, 51))
            else:
                # Degree / timeframe line (italic, muted gray)
                p = doc.add_paragraph()
                apply_paragraph_format(p, space_before=1, space_after=2, line_spacing=1.0)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(line)
                run.italic = True
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(85, 85, 85)

        # ── HIGHLIGHTS / CERTIFICATIONS / ACHIEVEMENTS ───────────────────────
        elif any(kw in current_section for kw in ("HIGHLIGHTS", "CERTIFICATIONS", "ACHIEVEMENTS")):
            if line.startswith('- '):
                add_body_bullet(doc, line[2:].strip())

        # ── ANY OTHER SECTION (e.g. PORTFOLIO, PROJECTS, PUBLICATIONS) ───────
        # Generic fallback so a section the CV happens to use is never silently
        # dropped. Renders with the same visual vocabulary as the named sections.
        else:
            if line.startswith('- '):
                add_body_bullet(doc, line[2:].strip())
            elif line.startswith('### '):
                add_subheader(doc, line[4:].strip(), size=10.5, space_before=8)
            elif line.startswith('**') and line.endswith('**') and len(line) > 4:
                add_subheader(doc, line[2:-2].strip())
            elif line.startswith('*') and line.endswith('*') and len(line) > 2:
                # Italic section lead-in
                add_italic_note(doc, line[1:-1].strip())
            else:
                p = doc.add_paragraph()
                apply_paragraph_format(p, space_before=2, space_after=4, line_spacing=1.15)
                add_formatted_text(p, line, font_size=9.5, color=RGBColor(51, 51, 51))

        i += 1

    doc.save(docx_path)
    print(f"Successfully generated: {docx_path}")


if __name__ == '__main__':
    main()
